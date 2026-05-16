"""Geração de documento Word (.docx) da OS usando python-docx."""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN  = RGBColor(0x16, 0xa3, 0x4a)
GREEN_D= RGBColor(0x15, 0x80, 0x3d)
DARK   = RGBColor(0x0f, 0x17, 0x2a)
MUTED  = RGBColor(0x64, 0x74, 0x8b)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _heading(doc, texto, level=1):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(11 if level == 1 else 9)
    run.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    # fundo indigo via shading no parágrafo — simplificado via tabela
    return p


def gerar_word_os(os_obj, upload_folder):
    nome    = f"{os_obj.numero_os}-{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.docx"
    caminho = os.path.join(upload_folder, nome)

    doc = Document()

    # Margens
    for sec in doc.sections:
        sec.left_margin   = Cm(1.5)
        sec.right_margin  = Cm(1.5)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)

    # ── Cabeçalho ────────────────────────────────────────
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    c0, c1 = table.rows[0].cells
    _set_cell_bg(c0, '15803d')
    _set_cell_bg(c1, '166534')

    r0 = c0.paragraphs[0].add_run('GERENCIADOR DE OS PARA CAMPO')
    r0.bold = True; r0.font.size = Pt(13); r0.font.color.rgb = WHITE

    r1 = c1.paragraphs[0].add_run(f'Nº {os_obj.numero_os}')
    r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = WHITE
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    def secao(titulo):
        t = doc.add_table(rows=1, cols=1)
        t.style = 'Table Grid'
        cell = t.rows[0].cells[0]
        _set_cell_bg(cell, '16a34a')
        run = cell.paragraphs[0].add_run(titulo)
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = WHITE
        doc.add_paragraph()

    def campo_valor(label, valor):
        p = doc.add_paragraph()
        lab = p.add_run(f'{label}: ')
        lab.bold = True
        lab.font.color.rgb = MUTED
        p.add_run(str(valor or '—'))
        p.paragraph_format.space_after = Pt(2)

    # ── Identificação ────────────────────────────────────
    secao('1. IDENTIFICAÇÃO DA OS')
    campo_valor('Número OS',            os_obj.numero_os)
    campo_valor('Status',               os_obj.status)
    campo_valor('Prioridade',           os_obj.prioridade)
    campo_valor('Tipo de Ocorrência',   os_obj.tipo_ocorrencia)
    campo_valor('Data Entrada',         os_obj.data_entrada)
    campo_valor('Data Saída',           os_obj.data_saida)
    campo_valor('Cliente',              os_obj.cliente.nome if os_obj.cliente else '—')
    campo_valor('Acompanhou a Execução', os_obj.acompanhante)  # RF12
    campo_valor('Local / Endereço',     os_obj.geo_endereco)
    campo_valor('Técnico Responsável',  os_obj.tecnico.usuario.nome if os_obj.tecnico else '—')
    doc.add_paragraph()

    # ── Laudo ────────────────────────────────────────────
    secao('2. LAUDO TÉCNICO & SOLUÇÃO')
    for label, val in [('Defeito Relatado', os_obj.defeito_relatado),
                        ('Condições Físicas', os_obj.condicoes_fisicas),
                        ('Laudo / Diagnóstico', os_obj.laudo_tecnico),
                        ('Solução Aplicada', os_obj.solucao_aplicada),
                        ('Peças Utilizadas', os_obj.pecas_utilizadas)]:
        campo_valor(label, val)
    doc.add_paragraph()

    # ── Checklist ────────────────────────────────────────
    if os_obj.checklist:
        secao('3. CHECKLIST DE TESTES')
        t = doc.add_table(rows=1, cols=4)
        t.style = 'Table Grid'
        headers = ['#', 'Teste', 'Feito', 'Técnico']
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            _set_cell_bg(cell, '15803d')
            r = cell.paragraphs[0].add_run(h)
            r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9)

        for idx, item in enumerate(os_obj.checklist, 1):
            row = t.add_row().cells
            row[0].text = str(idx)
            row[1].text = item.item_nome
            row[2].text = '✓' if item.feito else '✗'
            row[3].text = item.tecnico_verificador or ''
        doc.add_paragraph()

    # ── Termos ───────────────────────────────────────────
    if os_obj.termos_observacoes:
        secao('4. TERMOS / OBSERVAÇÕES')
        doc.add_paragraph(os_obj.termos_observacoes)
        doc.add_paragraph()

    # ── Assinaturas ──────────────────────────────────────
    secao('5. ASSINATURAS')
    sig_t = doc.add_table(rows=2, cols=2)
    sig_t.style = 'Table Grid'
    sig_t.rows[0].cells[0].text = 'Assinatura do Cliente'
    sig_t.rows[0].cells[1].text = 'Assinatura do Técnico'
    sig_t.rows[1].cells[0].text = '\n\n\n'
    sig_t.rows[1].cells[1].text = '\n\n\n'

    # ── Rodapé ───────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f'TECPOINT · Gerenciador OS para Campo · UniSENAI MT · Gerado em {datetime.now().strftime("%d/%m/%Y %H:%M")}'
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(7)
    footer_p.runs[0].font.color.rgb = MUTED

    doc.save(caminho)
    return caminho, nome
